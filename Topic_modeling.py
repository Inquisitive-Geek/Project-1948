import project_1948 as pjct
import nltk
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk import pos_tag
from nltk.tokenize import word_tokenize
from gensim import corpora, models
import docx as w
import matplotlib.pyplot as plt
import networkx as nx

#nltk.download('stopwords')
#nltk.download('averaged_perceptron_tagger')
#nltk.download('wordnet')
#pull interview files

#---RETRIEVING AND PREPARING INTERVIEWS---
print('Pulling interviews')
interview_1 = pjct.read_doc_file_to_string('Interview #1 - Participant 1.docx','rb')
interview_2 = pjct.read_doc_file_to_string('Interview #2-Participant 2.docx','rb')
interview_4 = pjct.read_doc_file_to_string('Interview#4Participant#4.docx','rb')
interview_5 = pjct.read_doc_file_to_string('Interview #5. Participant # 5.docx','rb')
interview_6 = pjct.read_doc_file_to_string('Interview #6. Participant #6.docx','rb')
interview_7 = pjct.read_doc_file_to_string('Interview #7. Participant #7..docx','rb')
interview_9 = pjct.read_doc_file_to_string('Interview #9. Participant #9.docx','rb')
interview_10 = pjct.read_doc_file_to_string('Interview #10. Participant #10.docx','rb')
interview_11 = pjct.read_doc_file_to_string('Interview #11 (1).docx','rb')
interview_12 = pjct.read_doc_file_to_string('Interview #12.docx','rb')
interview_13 = pjct.read_doc_file_to_string('Interview #13. Participant #13.docx','rb')
interview_14 = pjct.read_doc_file_to_string('interviews_new/Adnan.docx', 'rb')
interview_15 = pjct.read_doc_file_to_string('interviews_new/Amer.docx', 'rb')
interview_16 = pjct.read_doc_file_to_string('interviews_new/Aris.docx', 'rb')
interview_17 = pjct.read_doc_file_to_string('interviews_new/Belma.docx', 'rb')
interview_18 = pjct.read_doc_file_to_string('interviews_new/Danijel.docx', 'rb')
interview_19 = pjct.read_doc_file_to_string('interviews_new/Darko.docx', 'rb')
interview_20 = pjct.read_doc_file_to_string('interviews_new/Jusuf.docx', 'rb')
interview_21 = pjct.read_doc_file_to_string('interviews_new/Kenan.docx', 'rb')

interview_list = [interview_1, interview_2, interview_4, interview_5, interview_6, interview_7,
                  interview_9, interview_10, interview_11, interview_12, interview_13,
                  interview_14, interview_15]
interview_list2 = [interview_16,interview_17,interview_18,
                  interview_19,interview_20,interview_21]

#words to ignore
stop_words = set(nltk.corpus.stopwords.words('english'))
stop_words_add_on = {'okay','like','think','hello','feel','uhm','umm','really','something',
                     'yeah','say','want','yes','lot','make','way','well','get','ono','jel',
                     'ima','mhm','always','every','would','know','see','one','maybe','maybe',
                     'also','said','much','going','little','tell','even','day','thing','quite'
                     'gora','things','people','still','could','show'}
stop_words = stop_words | stop_words_add_on
stemmer = PorterStemmer()

#creating document collection
print('creating corpus...')
interview_tokens = []
j=0
for interview in interview_list:
    interview = pjct.remove_grouped_text(interview)
    interview = pjct.remove_punctuation_characters(interview,interviewer=' ',interviewee=' ')
    interview = interview.lower()
    interview_list[j]=interview
    j = j+1
    words = interview.split()
    words = [stemmer.stem(word).lower() for word in words if word.lower() not in stop_words]
    words = list({word for word in words if len(word)>2})
    interview_tokens.append(words)
print('corpus created')

#---IMPLEMENTING TOPIC MODELS---
    
#creating word-id dictionary and corpus (bag of words)
num_topics = 10
num_words = 30
dictionary = corpora.Dictionary(interview_tokens)
corpus = [dictionary.doc2bow(interview) for interview in interview_tokens]
del(interview)

#Tfidf model
tfidf = models.TfidfModel(corpus)
weighted_corpus = tfidf[corpus]

#LDA model implementation
print('---Begin LDA---\n')
LDA = models.ldamodel.LdaModel(weighted_corpus, num_topics = num_topics, id2word=dictionary, passes=30)
LDA.print_topics(num_topics=num_topics, num_words=num_words)
LDA_results = LDA.show_topics(num_topics=num_topics,num_words=num_words)
print('---End LDA---\n')

#HDP model implementation
print('---Begin HDP---\n')
HDP = models.hdpmodel.HdpModel(weighted_corpus,id2word=dictionary)
HDP.print_topics(num_topics=num_topics,num_words=num_words)
HDP_results = HDP.show_topics(num_topics=num_topics,num_words=num_words)
print('---End HDP---\n')

#LSI model implementation
print('---Begin LSI---\n')
LSI = models.lsimodel.LsiModel(weighted_corpus,num_topics=num_topics,id2word=dictionary)
LSI.print_topics(num_topics=num_topics,num_words=num_words)
LSI_results = LSI.show_topics(num_topics=num_topics,num_words=num_words)
print('---End LSI---\n')

#create document with results...4 topics, 30 wds each for each model.
'''folder = '/Users/JaredGallegos/Desktop/Project_1948/Interviews/'
results = w.Document()
results.add_heading('LDA model results',1)
for topic in LDA_results:
    results.add_paragraph(str(topic))
results.add_heading('HDP model results',1)
for topic in HDP_results:
    results.add_paragraph(str(topic))
results.add_heading('LSI model results',1)
for topic in HDP_results:
    results.add_paragraph(str(topic))
results.save(folder+'Topic_model_results.docx')
LDA.save('LDA_model')
HDP.save('HDP_model')
LSI.save('LSI_model')
print('Results and models saved')
'''
#---CREATING TOPIC GRAPHS---
topics_for_connection = 1

LDA_stats = {}
HDP_stats = {}
LSI_stats = {}

LDA_graph = nx.Graph()
HDP_graph = nx.Graph()
LSI_graph = nx.Graph()

interview_stats = [LDA_stats, HDP_stats,LSI_stats]
graphs = [LDA_graph, HDP_graph,LSI_graph]
models_list = [LDA, HDP, LSI]
model_indices = {0:'LDA Model',1:'HDP Model',2:'LSI Model'}
g=0
n=0

print('creating graphs')
#determine topics in each interview and add interview to graph
for model in models_list:
    for interview in interview_list2: 
        interview_tokens = interview.split()
        interview_bow = dictionary.doc2bow(interview_tokens)
        interview_stats[g][n] = model[interview_bow]
        graphs[g].add_node(n)
        n+=1
    g+=1
    n=0



#create edges between nodes that have a similar topic (topics only count if p>.50)
#one graph for each model
labels=[]
g=0
for graph in graphs:
    nodes_to_check = set(graph.nodes())
    dict_of_labels = {}
    for node in graphs[g].nodes():
        nodes_to_check = nodes_to_check - {node}#remove node from set of nodes to compare with
        node_topics = []
        topics=''
        for pair in interview_stats[g][node]: #get topics for the node.  Set labels.
            node_topics.append(pair[0])
            print('\t'+str(pair[0]))
            topics = topics + str(pair[0])
        print(topics)
        dict_of_labels[node] = topics
        print(dict_of_labels)
        for other_node in nodes_to_check: #check through for similar topics of nodes.  Add edge if similar topic.
            for pair in interview_stats[g][other_node]:
                if pair[0] in node_topics:
                    graph.add_edge(node,other_node)
                    break
    labels.append(dict_of_labels)
    g+=1

#Plot graph using shell layout.  Shells determined by number of topics contained in interviews
g=0
for graph in graphs:
    shells = []
    nodes = graph.nodes()
    nodes.sort(key = lambda node:len(interview_stats[g][node]),reverse=True)#sort by decreasing number of topics
    while len(nodes)!= 0:
        current_num_topics = len(interview_stats[g][nodes[0]])
        break_index = 0
        try:
            while len(interview_stats[g][nodes[break_index]])==current_num_topics:
                break_index+=1
        except IndexError:
            pass
        shells.append(nodes[:break_index])
        nodes = nodes[break_index:]
    plt.figure()
    nx.draw_networkx(graph,nx.shell_layout(graph,nlist=shells),labels=labels[g])
    plt.title(model_indices[g])
    g+=1
    
#Getting vertex degree information
print('---Begin graph stats---\n')
avg_deg = [0,0,0]
degrees = [[],[],[]]
g=0
for graph in graphs:
   avg_deg[g] = 2*len(graph.edges())/float(len(graph.nodes()))
   for node in graph.nodes():
       degrees[g].append(graph.degree(node))
   g+=1
print('Average degrees: '+str(avg_deg))
for j in range(len(degrees)):
    print('Graph '+str(j+1)+' degrees: '+str(degrees[j])+'\n')

#are graphs isomorphic?
print('LDA and HDP Isomorphic?: '+str(nx.is_isomorphic(LDA_graph,HDP_graph)))
print('LDA and LSI Isomorphic?: '+str(nx.is_isomorphic(LDA_graph,LSI_graph)))
print('HDP and LSI Isomorphic?: '+str(nx.is_isomorphic(HDP_graph,LSI_graph)))
print('---End graph stats---\n')